#!/usr/bin/env python3
"""Generate a Word document from the Release Management Plan markdown."""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, font_size=None, color=None, alignment=None, space_after=None, space_before=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "1e3a5f")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "f2f2f2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

def build_document():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default font --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ═══════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    add_formatted_paragraph(doc, 'Release Management Plan', bold=True, font_size=28,
                           color=(0x1e, 0x3a, 0x5f), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0x3b, 0x82, 0xf6)
    run.font.size = Pt(14)

    add_formatted_paragraph(doc, 'Enterprise Microservice CI/CD Deployment', bold=False, font_size=16,
                           color=(0x55, 0x55, 0x55), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'Recipe Management Application', bold=False, font_size=14,
                           color=(0x77, 0x77, 0x77), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    meta_items = [
        ('Document Version:', '2.0'),
        ('Date:', 'March 2026'),
        ('Module:', 'DevOps — Continuous Assessment'),
        ('Author:', 'DevOps Engineering Team'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(label + ' ')
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_l.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
        run_v = p.add_run(value)
        run_v.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, 'Table of Contents', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=12)

    toc_entries = [
        '1. Introduction',
        '2. System Architecture Overview',
        '3. Repository and Branching Strategy',
        '4. CI/CD Pipeline Architecture',
        '5. Deployment Strategy',
        '6. Infrastructure as Code',
        '7. Container Orchestration',
        '8. Evaluation Criteria',
        '    8.1 Performance',
        '    8.2 Ease of Configuration and Installation',
        '    8.3 Cost and Licensing',
        '    8.4 Monitoring and Logging',
        '    8.5 Scaling — Horizontal and Vertical',
        '    8.6 Backup and Restore Strategy',
        '    8.7 Security',
        '    8.8 Support',
        '    8.9 Vulnerability Checks on Images',
        '    8.10 Sustainability',
        '    8.11 Rollback Plan',
        '9. Change Management and Configuration Drift',
        '10. Version Unification',
        '11. Additional Features',
        '12. Conclusion and Critical Reflection',
        '13. References',
    ]
    for entry in toc_entries:
        indent = entry.startswith('    ')
        p = doc.add_paragraph()
        run = p.add_run(entry.strip())
        run.font.size = Pt(11 if not indent else 10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if indent:
            p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 1 — INTRODUCTION
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '1. Introduction', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_formatted_paragraph(doc, '1.1 Purpose', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    doc.add_paragraph(
        'This document constitutes the Release Management Plan (RMP) for an enterprise-style CI/CD deployment '
        'of a microservice-based recipe management application. The RMP serves as both a specification of the '
        'deployment architecture and a critical evaluation of the decisions taken, assessed against industry-standard criteria.'
    )
    doc.add_paragraph(
        'The transition from monolithic to microservice architectures represents a fundamental shift in how software '
        'is built, deployed, and operated (Newman, 2021). Monolithic applications — where user-facing logic, business '
        'rules, and data access are tightly coupled into a single deployable unit — present significant challenges for '
        'independent scaling, fault isolation, and deployment velocity. The microservice architectural style addresses '
        'these concerns by decomposing applications into small, independently deployable services that communicate over '
        'well-defined APIs (Fowler and Lewis, 2014).'
    )
    doc.add_paragraph(
        'However, the microservice approach introduces its own complexities: distributed system concerns, inter-service '
        'communication, data consistency across service boundaries, and — critically for this report — significantly more '
        'complex deployment orchestration. Where a monolith requires a single deployment pipeline, a microservice '
        'architecture demands coordinated but independent deployment of multiple services, each with its own lifecycle.'
    )

    add_formatted_paragraph(doc, '1.2 Scope', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph('This RMP covers:')
    scope_items = [
        'The architectural design of a three-component microservice application (frontend, backend API, database).',
        'The fully automated CI/CD pipeline from code commit to production deployment.',
        'Infrastructure provisioning via Infrastructure as Code (IaC).',
        'The deployment strategy (blue/green) and its implementation on Kubernetes.',
        'A critical evaluation against eleven criteria: performance, ease of configuration, cost, monitoring, scaling, backup/restore, security, support, vulnerability scanning, sustainability, and rollback.',
    ]
    for item in scope_items:
        p = doc.add_paragraph(style='List Bullet')
        p.text = item

    add_formatted_paragraph(doc, '1.3 Design Philosophy', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph('Three principles guided every decision in this plan:')
    principles = [
        ('Automation over documentation', 'If a process can be automated, it must be. Manual runbooks create opportunities for human error and configuration drift. As Humble and Farley (2010) argue, "if it hurts, do it more frequently, and bring the pain forward."'),
        ('Tool choice follows requirements', 'Tools were selected to serve the deployment requirements, not the reverse. Each tool selection is justified against alternatives in the relevant section.'),
        ('Immutability and reproducibility', 'Infrastructure and application artefacts are treated as immutable. Changes are made by replacing, not modifying. This ensures that the system can be destroyed and reliably rebuilt at any time.'),
    ]
    for i, (title, desc) in enumerate(principles, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title} — ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 2 — SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '2. System Architecture Overview', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_formatted_paragraph(doc, '2.1 Application Components', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    doc.add_paragraph(
        'The application is a recipe management system composed of three loosely coupled services:'
    )

    add_table(doc,
        ['Service', 'Technology Stack', 'Responsibility', 'Port'],
        [
            ['Frontend (FE)', 'Node.js 20, Express 4, EJS', 'Dynamic web UI — renders HTML, proxies API requests', '3000'],
            ['Backend (BE)', 'Java 17, Spring Boot 3, Spring Data MongoDB', 'RESTful API — CRUD operations, validation, business logic', '8080'],
            ['Database', 'MongoDB 7', 'Persistent data storage — recipe documents as BSON', '27017'],
        ])

    doc.add_paragraph(
        'Each service is independently buildable and deployable. The frontend can be updated without touching the '
        'backend, and vice versa. This is the defining characteristic of a microservice architecture — independent '
        'deployability (Newman, 2021).'
    )

    add_formatted_paragraph(doc, '2.2 Communication Patterns', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph(
        'Communication is synchronous request-response via HTTP/JSON between FE and BE, and via the MongoDB wire '
        'protocol between BE and the database. This is the simplest viable integration pattern for this application\'s requirements.'
    )
    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Synchronous communication creates temporal coupling — if the backend is unavailable, the frontend cannot '
        'serve dynamic content. For a larger system, asynchronous messaging (e.g., via Azure Service Bus or RabbitMQ) '
        'would improve resilience. However, for a two-service application with simple CRUD operations, the added '
        'complexity of an event-driven architecture is not warranted.'
    )

    add_formatted_paragraph(doc, '2.3 Infrastructure Components', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Component', 'Chosen Tool', 'Purpose', 'Alternatives Considered'],
        [
            ['Cloud Provider', 'Microsoft Azure', 'Hosting all infrastructure', 'AWS, GCP'],
            ['Container Registry', 'Azure Container Registry (ACR)', 'Storing and serving Docker images', 'Docker Hub, GitHub Container Registry'],
            ['Container Orchestration', 'Azure Kubernetes Service (AKS)', 'Running containerised workloads', 'Azure Container Apps, Docker Swarm'],
            ['Infrastructure as Code', 'Terraform (HashiCorp)', 'Declarative provisioning', 'Bicep, Pulumi, ARM Templates'],
            ['Application Packaging', 'Helm 3', 'Templated K8s manifest management', 'Kustomize, raw manifests'],
            ['CI/CD Platform', 'GitHub Actions', 'Automated build, test, deploy', 'Jenkins, GitLab CI, Azure DevOps'],
            ['Monitoring', 'Azure Monitor + Log Analytics', 'Metrics, logs, alerting', 'Prometheus + Grafana, Datadog'],
            ['Backup Storage', 'Azure Blob Storage', 'MongoDB backup destination', 'Azure Files, AWS S3'],
        ])

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 3 — REPO & BRANCHING
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '3. Repository and Branching Strategy', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_formatted_paragraph(doc, '3.1 Mono-Repo vs Poly-Repo Evaluation', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    doc.add_paragraph(
        'A foundational decision in microservice deployment is whether services should share a single repository '
        '(mono-repo) or each reside in their own repository (poly-repo).'
    )

    add_table(doc,
        ['Criterion', 'Mono-Repo', 'Poly-Repo'],
        [
            ['Atomic cross-service changes', 'Single commit updates both FE and BE', 'Requires coordinated PRs across repos'],
            ['CI/CD simplicity', 'One repo; path-based triggers isolate pipelines', 'Each repo has own pipeline'],
            ['Access control', 'Coarser — all developers see all code', 'Fine-grained per-service permissions'],
            ['Dependency management', 'Shared infra code lives alongside services', 'Infra code needs own repo or duplication'],
            ['Build times', 'Can become slow; mitigated by path filters', 'Naturally isolated'],
            ['Onboarding', 'Single clone; complete project visibility', 'Multiple clones; harder to understand full system'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Decision: Mono-repo. ')
    run.bold = True
    p.add_run(
        'For this two-service application, a mono-repo provides the strongest developer experience. The decision '
        'would be revisited if the application grew to ten or more services. Google, Meta, and Microsoft all operate '
        'at mono-repo scale, demonstrating that tooling can mitigate scaling concerns (Potvin and Levenberg, 2016). '
        'The key enabler is path-based CI triggers in GitHub Actions.'
    )

    add_formatted_paragraph(doc, '3.2 Branching Model', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph('The project follows a GitHub Flow variant with an integration branch:')

    add_table(doc,
        ['Branch', 'Purpose', 'Protection Rules'],
        [
            ['main', 'Production-ready code; every commit is deployable', 'Requires CI pass, vuln scan, 1 approval'],
            ['develop', 'Integration branch for feature work', 'Requires passing CI'],
            ['feature/*', 'Individual feature development', 'No protection'],
            ['hotfix/*', 'Critical production fixes', 'Same as main'],
        ])

    doc.add_paragraph('Pull requests into main enforce: (1) all CI checks pass, (2) Trivy scan reports no CRITICAL CVEs, (3) at least one code review approval.')

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'GitHub Flow is simpler than GitFlow and suits continuous deployment well. However, it assumes that main '
        'is always deployable, which requires strong CI discipline. For teams with less CI maturity, GitFlow\'s '
        'explicit release branches provide an additional safety net.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 4 — CI/CD PIPELINE
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '4. CI/CD Pipeline Architecture', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_formatted_paragraph(doc, '4.1 Continuous Integration Design', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    doc.add_paragraph(
        'Each service has its own independent CI pipeline, triggered only when code in its directory changes. '
        'This achieves a key microservice goal: independent deployability.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Frontend CI Pipeline (ci-frontend.yml):')
    run.bold = True
    fe_steps = [
        'Checkout code', 'Setup Node.js 20', 'Install dependencies (npm install)',
        'Lint (ESLint)', 'Unit tests (Jest with coverage)', 'Build Docker image (multi-stage)',
        'Vulnerability scan (Trivy — fail on CRITICAL)', 'Push image to ACR (SHA tag + latest)',
    ]
    for i, step in enumerate(fe_steps, 1):
        p = doc.add_paragraph(style='List Bullet')
        p.text = f'{i}. {step}'
        p.paragraph_format.left_indent = Cm(1)

    p = doc.add_paragraph()
    run = p.add_run('Backend CI Pipeline (ci-backend.yml):')
    run.bold = True
    be_steps = [
        'Checkout code', 'Setup Java 17 (Temurin) + Maven cache',
        'Compile and run unit tests (mvn clean verify)', 'Upload test reports as artifacts',
        'Build Docker image (multi-stage)', 'Vulnerability scan (Trivy — fail on CRITICAL)',
        'Push image to ACR (SHA tag + latest)',
    ]
    for i, step in enumerate(be_steps, 1):
        p = doc.add_paragraph(style='List Bullet')
        p.text = f'{i}. {step}'
        p.paragraph_format.left_indent = Cm(1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Key design decisions:')
    run.bold = True

    decisions = [
        'Image tags use Git SHA prefixes (e.g., a1b2c3d) rather than version numbers — every deployed image maps back to an exact commit.',
        'Images are only built and pushed on main branch. Feature branches run tests but do not produce artefacts.',
        'Each pipeline is self-contained — it does not depend on or trigger other pipelines.',
    ]
    for d in decisions:
        p = doc.add_paragraph(style='List Bullet')
        p.text = d

    add_formatted_paragraph(doc, '4.2 Continuous Delivery Design', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph(
        'The CD pipeline (cd-deploy.yml) orchestrates deployment of all services to AKS. It is triggered '
        'automatically when either CI pipeline completes on main, or manually via workflow_dispatch.'
    )

    cd_steps = [
        'Azure Login (service principal credentials)',
        'Get AKS credentials (kubectl context)',
        'Create/verify namespace',
        'Determine target slot (auto-detect inactive slot)',
        'Deploy MongoDB (Helm upgrade --install)',
        'Deploy Backend to target slot (Helm)',
        'Deploy Frontend to target slot (Helm)',
        'Health check (kubectl wait for pod readiness)',
        'Switch traffic (patch service selectors)',
        'Print deployment summary',
    ]
    for i, step in enumerate(cd_steps, 1):
        p = doc.add_paragraph(style='List Bullet')
        p.text = f'{i}. {step}'
        p.paragraph_format.left_indent = Cm(1)

    add_formatted_paragraph(doc, '4.3 CI/CD Platform Selection', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Criterion', 'GitHub Actions', 'Jenkins', 'GitLab CI/CD', 'Azure DevOps'],
        [
            ['Setup complexity', 'Zero — built into GitHub', 'High — self-hosted', 'Moderate', 'Moderate'],
            ['Config language', 'YAML (declarative)', 'Groovy (imperative)', 'YAML', 'YAML'],
            ['Azure integration', 'Excellent — first-party actions', 'Via plugins', 'Manual setup', 'Native — best'],
            ['Secret management', 'Built-in encrypted secrets', 'Credentials plugin', 'CI/CD variables', 'Variable groups'],
            ['Cost (OSS)', '2,000 min/month free', 'Free (self-hosted cost)', '400 min/month free', '1,800 min/month free'],
            ['Learning curve', 'Low', 'Moderate-High', 'Low-Moderate', 'Low-Moderate'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Decision: GitHub Actions. ')
    run.bold = True
    p.add_run(
        'The mono-repo already lives on GitHub, making Actions the zero-friction choice. Jenkins was rejected for '
        'its operational overhead. Azure DevOps offers superior Azure-native integration but would require moving the repository.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'GitHub Actions\' main limitation is debugging — failed workflows require inspecting logs in the browser '
        'with no interactive debugging. Jenkins\' Blue Ocean and pipeline replay are superior for troubleshooting.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 5 — DEPLOYMENT STRATEGY
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '5. Deployment Strategy', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_formatted_paragraph(doc, '5.1 Evaluation of Deployment Strategies', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    add_table(doc,
        ['Strategy', 'Downtime', 'Rollback Speed', 'Resource Cost', 'Complexity'],
        [
            ['Recreate', 'Yes', 'Slow (full redeploy)', 'Low', 'Very Low'],
            ['Rolling Update', 'No', 'Moderate', 'Low', 'Low'],
            ['Blue/Green', 'No', 'Instant (switch back)', 'High (double)', 'Moderate'],
            ['Canary', 'No', 'Fast', 'Moderate', 'High (needs mesh)'],
        ])

    doc.add_paragraph(
        'Recreate was rejected due to unacceptable downtime. Rolling Update\'s slower rollback and mixed-version '
        'serving make it less suitable. Canary requires a service mesh, adding disproportionate complexity. '
        'Blue/Green was selected as the optimal balance of safety, simplicity, and rollback speed.'
    )

    add_formatted_paragraph(doc, '5.2 Blue/Green Implementation', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph(
        'The blue/green strategy is implemented at the Kubernetes level using pod labels and service selector '
        'patching. Each deployment is labelled with a slot label (blue or green). The Kubernetes Service\'s selector '
        'determines which slot receives traffic.'
    )

    doc.add_paragraph('Deploying a new version means:')
    bg_steps = [
        'Identifying the inactive slot (the one not currently receiving traffic).',
        'Deploying the new version\'s pods with the inactive slot\'s label.',
        'Waiting for all new pods to pass health checks.',
        'Patching the Service selector to point to the newly deployed slot.',
    ]
    for i, step in enumerate(bg_steps, 1):
        p = doc.add_paragraph(style='List Bullet')
        p.text = f'{i}. {step}'

    doc.add_paragraph(
        'The previous version\'s pods remain running on the now-inactive slot, enabling instant rollback by '
        'simply patching the selector back.'
    )

    add_formatted_paragraph(doc, '5.3 Traffic Switching Mechanism', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph(
        'Traffic switching is performed via kubectl patch — an atomic operation from the user\'s perspective. '
        'Kubernetes immediately begins routing new connections to the target pods. Existing connections to old pods '
        'are drained gracefully.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'This approach switches frontend and backend independently, creating a brief window where versions may be '
        'mismatched. For this application with a stable API contract, this is acceptable. In systems with strict '
        'version coupling, an ingress selector should be switched atomically.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 6 — INFRASTRUCTURE AS CODE
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '6. Infrastructure as Code', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_formatted_paragraph(doc, '6.1 IaC Tool Selection', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    add_table(doc,
        ['Criterion', 'Terraform', 'Bicep', 'Pulumi', 'ARM Templates'],
        [
            ['Cloud support', 'Multi-cloud (3000+ providers)', 'Azure only', 'Multi-cloud', 'Azure only'],
            ['Language', 'HCL (declarative)', 'Bicep DSL', 'Python/TS/Go/C#', 'JSON (verbose)'],
            ['State management', 'Explicit state file', 'Managed by Azure', 'Built-in state', 'Managed by Azure'],
            ['Plan/preview', 'terraform plan', 'what-if', 'pulumi preview', 'what-if'],
            ['Community', 'Largest ecosystem', 'Growing (Microsoft)', 'Moderate, growing', 'Large but declining'],
            ['Maturity', 'Very mature (2014)', 'Mature (2020)', 'Maturing (2018)', 'Very mature but legacy'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Decision: Terraform. ')
    run.bold = True
    p.add_run(
        'Three factors: (1) multi-cloud portability, (2) explicit state management enabling reliable plan and destroy, '
        '(3) unmatched community and ecosystem.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Bicep would be strong for Azure-only. Its transpilation to ARM means it\'s always up-to-date with Azure\'s '
        'latest features, whereas the Terraform AzureRM provider sometimes lags.'
    )

    add_formatted_paragraph(doc, '6.2 Terraform Architecture', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Resource', 'Terraform Type', 'Configuration'],
        [
            ['Resource Group', 'azurerm_resource_group', 'Naming: rg-{project}-{env}'],
            ['Virtual Network', 'azurerm_virtual_network', '10.0.0.0/16 address space'],
            ['AKS Subnet', 'azurerm_subnet', '10.0.1.0/24'],
            ['Container Registry', 'azurerm_container_registry', 'Basic SKU, admin disabled'],
            ['AKS Cluster', 'azurerm_kubernetes_cluster', '1–5 nodes, autoscaling, Azure CNI, Calico'],
            ['ACR Pull Role', 'azurerm_role_assignment', 'AKS kubelet identity → AcrPull'],
            ['Log Analytics', 'azurerm_log_analytics_workspace', '30-day retention, PerGB2018'],
            ['Storage Account', 'azurerm_storage_account', 'LRS, TLS 1.2, for backups'],
            ['Blob Container', 'azurerm_storage_container', 'Private access, mongo-backups'],
        ])

    add_formatted_paragraph(doc, '6.3 State Management and Drift Prevention', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph(
        'The project supports local state (dev) and remote state in Azure Blob Storage (teams). Configuration '
        'drift is prevented by making all changes through Terraform only — terraform plan produces a diff for '
        'review, and apply is the only authorised modification method.'
    )

    add_formatted_paragraph(doc, '6.4 Destroy and Rebuild Capability', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph(
        'terraform destroy removes all infrastructure. terraform apply re-creates it identically. '
        'scripts/deploy.sh rebuilds images and deploys to the new cluster. Total rebuild time: ~15–20 minutes. '
        'This validates IaC completeness, enables cost management (destroy non-prod overnight), and eliminates drift.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 7 — CONTAINER ORCHESTRATION
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '7. Container Orchestration', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_formatted_paragraph(doc, '7.1 Kubernetes on AKS', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    add_table(doc,
        ['Setting', 'Value', 'Justification'],
        [
            ['Kubernetes version', '1.29', 'Latest stable; security patches'],
            ['Node VM size', 'Standard_B2s', 'Burstable, cost-effective (2 vCPUs, 4 GB)'],
            ['Node count', '2 (autoscale 1–5)', 'HA baseline; autoscaler handles spikes'],
            ['Network plugin', 'Azure CNI', 'Pod IPs from VNet; required for network policies'],
            ['Network policy', 'Calico', 'Industry-standard pod-to-pod traffic control'],
            ['Identity', 'System-assigned MI', 'No service principal secrets to rotate'],
            ['Monitoring', 'OMS Agent → Log Analytics', 'Container logs/metrics to Azure Monitor'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'For production: add Availability Zones, dedicated node pools, and AKS Standard tier (99.95% SLA).'
    )

    add_formatted_paragraph(doc, '7.2 Helm-Based Application Packaging', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Chart', 'Key Resources', 'Notes'],
        [
            ['frontend/', 'Deployment, Service (LoadBalancer), HPA', 'Blue/green via deployment.slot value'],
            ['backend/', 'Deployment, Service (ClusterIP), HPA', 'Internal-only service'],
            ['mongodb/', 'StatefulSet, Service, PVC, CronJob', 'Persistent storage, automated backup'],
        ])

    doc.add_paragraph(
        'Why Helm over Kustomize: templating (Go templates for parameterised manifests), release management '
        '(tracks history, enables helm rollback), lifecycle hooks, and extensive chart ecosystem.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 8 — EVALUATION CRITERIA
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '8. Evaluation Criteria', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    # 8.1 Performance
    add_formatted_paragraph(doc, '8.1 Performance', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    add_table(doc,
        ['Measure', 'Implementation', 'Impact'],
        [
            ['Minimal container images', 'Multi-stage builds: FE ~80 MB, BE ~150 MB', 'Faster pod startup'],
            ['Readiness probes', 'FE: /health; BE: /actuator/health', 'Traffic only to ready pods'],
            ['Connection pooling', 'MongoDB driver pool (100 max); Axios keep-alive', 'Avoids per-request overhead'],
            ['Application metrics', 'Prometheus histograms on both FE and BE', 'Continuous performance monitoring'],
            ['Resource requests/limits', 'CPU and memory set per pod', 'Predictable performance'],
        ])

    add_table(doc,
        ['Metric', 'Frontend', 'Backend'],
        [
            ['Cold start time', '~1s (Node.js)', '~15–30s (JVM + Spring)'],
            ['Warm request latency', '<50ms', '<20ms'],
            ['Throughput ceiling (single pod)', '~500 req/s', '~200 req/s'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'JVM cold start is the primary concern. Spring Boot 3 AOT or GraalVM native image could reduce startup to '
        '<1s at the cost of build complexity. Mitigated by readiness probes and min 2 replicas. Spring WebFlux or '
        'response caching would improve throughput but add disproportionate complexity.'
    )

    # 8.2 Ease of Configuration
    add_formatted_paragraph(doc, '8.2 Ease of Configuration and Installation', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph('Local development: a single docker-compose up --build starts all services. No local Java, Node.js, or MongoDB required.')

    add_table(doc,
        ['Step', 'Command', 'Time'],
        [
            ['Clone repository', 'git clone <repo-url>', '30s'],
            ['Start all services', 'docker-compose up --build', '3–5 min'],
            ['Access application', 'http://localhost:3000', '—'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Cloud setup requires four CLI tools (az, terraform, helm, kubectl). GitHub Codespaces with a devcontainer '
        'would eliminate local tooling requirements. Azure Container Apps would simplify orchestration but sacrifice '
        'fine-grained control.'
    )

    # 8.3 Cost
    add_formatted_paragraph(doc, '8.3 Cost and Licensing', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Component', 'Dev Cost/Month', 'Prod Cost/Month'],
        [
            ['AKS Control Plane', '$0 (free tier)', '~$73 (Standard)'],
            ['AKS Node VMs', '~$60 (2× B2s)', '~$210 (3× D2s_v3)'],
            ['ACR', '~$5 (Basic)', '~$20 (Standard)'],
            ['Log Analytics', '~$10', '~$35'],
            ['Storage', '<$1', '<$5'],
            ['TOTAL', '~$75/month', '~$345/month'],
        ])

    doc.add_paragraph(
        'All tools are open-source or free-tier. Cost optimisation: scheduled teardown (~65% savings), '
        'autoscaler tuning (min 1 node), spot instances (up to 90% discount), reserved instances (30–60% savings).'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Azure Container Apps (consumption pricing) could reduce costs to $10–20/month for low-traffic apps. '
        'MongoDB Atlas free tier (M0) would eliminate in-cluster database costs.'
    )

    # 8.4 Monitoring
    add_formatted_paragraph(doc, '8.4 Monitoring and Logging', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph(
        'All services expose Prometheus-format metrics. The AKS OMS agent collects container logs and metrics, '
        'forwarding them to Azure Log Analytics. Services log to stdout/stderr following twelve-factor principles.'
    )

    add_table(doc,
        ['Alert', 'Condition', 'Severity'],
        [
            ['Pod restart loop', 'Restarts > 3 in 5 min', 'Critical'],
            ['High error rate', 'HTTP 5xx > 5% for 2 min', 'Warning'],
            ['High latency', 'P95 > 2s for 5 min', 'Warning'],
            ['Node pressure', 'CPU > 90% for 10 min', 'Warning'],
            ['MongoDB failure', 'Connection errors > 0', 'Critical'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Prometheus + Grafana would provide richer dashboards. OpenTelemetry distributed tracing should be added '
        'for cross-service request correlation.'
    )

    # 8.5 Scaling
    add_formatted_paragraph(doc, '8.5 Scaling — Horizontal and Vertical', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Service', 'Min Replicas', 'Max Replicas', 'CPU Target', 'Memory Target'],
        [
            ['Frontend', '2', '10', '70%', '80%'],
            ['Backend', '2', '8', '70%', '80%'],
            ['MongoDB', '1', '1', 'N/A', 'N/A'],
        ])

    doc.add_paragraph(
        'The AKS cluster autoscaler manages node count (1–5). Scale out when pods are pending; '
        'scale in after 10 min idle.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'MongoDB is a single instance — production needs a Replica Set or Azure Cosmos DB for MongoDB API. '
        'HPA is reactive; KEDA could enable proactive scaling based on external metrics.'
    )

    # 8.6 Backup
    add_formatted_paragraph(doc, '8.6 Backup and Restore Strategy', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Metric', 'Target', 'Justification'],
        [
            ['RPO (max data loss)', '24 hours', 'Daily backups; recipe data changes infrequently'],
            ['RTO (max recovery time)', '1 hour', 'Time to restore from backup and redeploy'],
        ])

    doc.add_paragraph(
        'A Kubernetes CronJob runs mongodump daily at 02:00 UTC, compresses with tar/gzip, and maintains '
        '7-day rolling retention.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Backups should be uploaded to Azure Blob Storage for durability. Monthly restore tests, encryption at rest, '
        'and point-in-time recovery via oplog should be implemented. Cosmos DB provides built-in PITR.'
    )

    # 8.7 Security
    add_formatted_paragraph(doc, '8.7 Security', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Layer', 'Measure', 'Implementation'],
        [
            ['Network', 'VNet isolation', 'AKS on dedicated subnet (10.0.1.0/24)'],
            ['Network', 'Network policies', 'Calico for pod-to-pod traffic control'],
            ['Container', 'Non-root execution', 'Dockerfiles create appuser; runAsNonRoot: true'],
            ['Container', 'Minimal images', 'Alpine-based (~5 MB base)'],
            ['Container', 'Multi-stage builds', 'Source code excluded from runtime images'],
            ['Registry', 'No admin credentials', 'Managed identity with AcrPull role'],
            ['Application', 'Security headers', 'Helmet.js on frontend'],
            ['Application', 'Input validation', 'Jakarta Bean Validation on backend'],
            ['Application', 'Rate limiting', 'express-rate-limit (100 req/15 min)'],
            ['Supply chain', 'Vulnerability scanning', 'Trivy in CI — fail on CRITICAL'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Production needs: explicit NetworkPolicy resources, Pod Security Standards (Restricted), Azure Key Vault '
        'for secrets, Application Gateway WAF, cert-manager + Let\'s Encrypt, and image signing (Notary v2).'
    )

    # 8.8 Support
    add_formatted_paragraph(doc, '8.8 Support', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Tool', 'Support Model', 'Critical Response Time'],
        [
            ['AKS / ACR', 'Microsoft Unified Support', '15 min – 1 hour'],
            ['Terraform', 'Community / Enterprise', 'Best-effort / 1 hour'],
            ['Helm', 'CNCF Community', 'Best-effort'],
            ['GitHub Actions', 'GitHub Support', '30 min (Premium)'],
            ['MongoDB Community', 'Community forums', 'Best-effort'],
            ['Spring Boot', 'Community / Tanzu', 'Best-effort / contract'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'MongoDB Community has no vendor SLA. For mission-critical data, MongoDB Atlas (99.995% SLA) or '
        'Azure Cosmos DB for MongoDB API would eliminate this gap.'
    )

    # 8.9 Vulnerability
    add_formatted_paragraph(doc, '8.9 Vulnerability Checks on Images', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Feature', 'Trivy', 'Docker Scout', 'Snyk Container'],
        [
            ['Cost', 'Free (Apache 2.0)', 'Free tier (limited)', 'Free tier (limited)'],
            ['OS packages', 'Yes', 'Yes', 'Yes'],
            ['Language deps', 'Yes', 'Yes', 'Yes'],
            ['IaC scanning', 'Yes', 'No', 'Yes'],
            ['License scanning', 'Yes', 'No', 'Yes'],
        ])

    doc.add_paragraph(
        'Severity policy: CRITICAL = build fails; HIGH = warning; MEDIUM/LOW = informational.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Trivy scans only at build time. ACR Defender provides continuous scanning. Dependabot automates dependency updates.'
    )

    # 8.10 Sustainability
    add_formatted_paragraph(doc, '8.10 Sustainability', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Measure', 'Implementation', 'Impact'],
        [
            ['Right-sizing', 'B-series burstable VMs', 'Lower baseline energy'],
            ['Autoscaling', 'HPA + cluster autoscaler', 'Resources only when needed'],
            ['Environment teardown', 'Destroy non-prod overnight', '~65% compute savings'],
            ['Minimal images', 'Alpine-based Docker images', 'Less storage and transfer'],
            ['Region selection', 'Azure renewable-energy regions', 'Lower carbon footprint'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Most impactful measure is automating environment teardown. Serverless architecture (Azure Functions + '
        'Cosmos DB) would eliminate baseline consumption entirely.'
    )

    # 8.11 Rollback
    add_formatted_paragraph(doc, '8.11 Rollback Plan', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    add_table(doc,
        ['Tier', 'Mechanism', 'Recovery Time', 'When to Use'],
        [
            ['Tier 1', 'Slot switch (kubectl patch)', '<10 seconds', 'App bug post-deployment'],
            ['Tier 2', 'Helm rollback', '<60 seconds', 'Config change caused issues'],
            ['Tier 3', 'Previous image tag redeploy', '<5 minutes', 'Need to go back multiple versions'],
            ['Tier 4', 'Infrastructure rebuild (terraform)', '15–30 minutes', 'Infrastructure compromise'],
        ])

    doc.add_paragraph(
        'Tier 1 is instant because both blue and green slots maintain running pods. Switching the service selector '
        'is a metadata-only operation — no image pulls, scheduling, or health check waits.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'Current rollback is manual. Production should automate: monitor health metrics post-deploy and auto-revert '
        'if error rates exceed thresholds. Data failures require database restore (Tier 4).'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 9 — CHANGE MANAGEMENT
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '9. Change Management and Configuration Drift', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    doc.add_paragraph(
        'Configuration drift occurs when actual infrastructure state diverges from intended state, typically '
        'due to manual changes. Morris (2016) identifies drift as the primary obstacle to Continuous Delivery.'
    )

    add_table(doc,
        ['Configuration Type', 'Source of Truth', 'Applied By'],
        [
            ['Azure resources', 'Terraform files', 'terraform apply'],
            ['Kubernetes workloads', 'Helm charts', 'helm upgrade --install'],
            ['Application config', 'Helm values.yaml + --set overrides', 'Helm'],
            ['CI/CD workflows', 'GitHub Actions YAML', 'GitHub'],
            ['Container contents', 'Dockerfiles', 'docker build'],
            ['Secrets', 'GitHub encrypted secrets', 'GitHub Actions'],
        ])

    doc.add_paragraph(
        'No manual changes are permitted. terraform apply and helm upgrade are idempotent — they correct drift '
        'back to the desired state. Every change follows code → PR → CI → merge → CD → production.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 10 — VERSION UNIFICATION
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '10. Version Unification', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_table(doc,
        ['Component', 'Version Pinning', 'Location'],
        [
            ['Node.js', 'node:20-alpine image tag', 'Dockerfile'],
            ['Java', 'java.version=17; temurin:17-jre-alpine', 'pom.xml, Dockerfile'],
            ['Spring Boot', 'spring-boot-starter-parent:3.3.2', 'pom.xml'],
            ['Terraform', 'required_version >= 1.5.0', 'providers.tf'],
            ['AzureRM provider', '~> 3.100', 'providers.tf'],
            ['Kubernetes', '1.29', 'variables.tf'],
            ['MongoDB', 'mongo:7 image tag', 'Helm values, docker-compose'],
            ['GitHub Actions', 'Pinned action versions (@v4, @v5)', 'workflow YAML'],
        ])

    p = doc.add_paragraph()
    run = p.add_run('Critique: ')
    run.bold = True
    run.italic = True
    p.add_run(
        'npm\'s package-lock.json should be committed and npm ci used instead of npm install for fully '
        'reproducible builds.'
    )

    # ═══════════════════════════════════════════
    # SECTION 11 — ADDITIONAL FEATURES
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '11. Additional Features', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8, space_before=16)

    features = [
        ('Automated MongoDB Backup (CronJob)', 'Kubernetes CronJob runs mongodump daily at 02:00 UTC with 7-day retention and auto-cleanup.'),
        ('Horizontal Pod Autoscaling (HPA)', 'HPA v2 with multi-metric scaling (CPU + memory) on frontend (2–10 pods) and backend (2–8 pods).'),
        ('Container Vulnerability Scanning (Trivy)', 'Integrated into CI — scans OS packages + language dependencies. Builds fail on CRITICAL CVEs.'),
        ('Blue/Green Zero-Downtime Deployment', 'Pod label slot switching with automatic inactive-slot detection and instant rollback (<10s).'),
    ]
    for i, (title, desc) in enumerate(features, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title}')
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 12 — CONCLUSION
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '12. Conclusion and Critical Reflection', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_formatted_paragraph(doc, '12.1 Summary of Achievements', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6)

    add_table(doc,
        ['Requirement', 'Implementation'],
        [
            ['2+ communicating services', 'Frontend (Node.js) ↔ Backend (Spring Boot) via HTTP/JSON'],
            ['Data layer', 'MongoDB 7 with persistent storage and automated backup'],
            ['Fully automated CI/CD', 'GitHub Actions with per-service CI + unified CD'],
            ['Infrastructure as Code', 'Terraform provisions AKS, ACR, VNet, monitoring, storage'],
            ['Automated release management', 'Helm-based deployment with blue/green slot switching'],
            ['Destroy and rebuild', 'terraform destroy + apply rebuilds in ~15 minutes'],
            ['Consistent change management', 'All changes via code → PR → CI → merge → CD'],
            ['Additional features (4)', 'Backup CronJob, HPA, Trivy scanning, blue/green deployment'],
        ])

    add_formatted_paragraph(doc, '12.2 Strengths', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)
    strengths = [
        'Complete automation — zero manual steps from commit to production.',
        'Independent service pipelines — per-service CI with path filters.',
        'Robust rollback — blue/green with sub-10-second slot switching.',
        'Security by default — non-root, minimal images, vulnerability scanning, managed identity.',
        'Observable system — Prometheus metrics, structured logging, Azure Monitor.',
    ]
    for s in strengths:
        p = doc.add_paragraph(style='List Bullet')
        p.text = s

    add_formatted_paragraph(doc, '12.3 Limitations and Future Work', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)
    limitations = [
        'Single MongoDB instance — production needs Replica Set or Cosmos DB.',
        'No integration or end-to-end tests — add FE→BE→DB test stage.',
        'No service mesh — needed for canary deployments, mutual TLS, tracing.',
        'No GitOps — ArgoCD/Flux for declarative cluster state reconciliation.',
        'Manual rollback — automate with post-deploy health gates.',
        'No authentication — add Azure AD/Entra ID with OAuth 2.0/OIDC.',
    ]
    for l in limitations:
        p = doc.add_paragraph(style='List Bullet')
        p.text = l

    add_formatted_paragraph(doc, '12.4 Concluding Remarks', bold=True, font_size=14,
                           color=(0x2d, 0x5a, 0x8a), space_after=6, space_before=12)

    doc.add_paragraph(
        'The microservice architectural style introduces deployment complexity that monolithic applications do not face. '
        'This RMP demonstrates that with appropriate tooling — Terraform for infrastructure, Helm for application '
        'packaging, and GitHub Actions for pipeline orchestration — this complexity can be fully managed through automation.'
    )
    doc.add_paragraph(
        'The result is a deployment system that is reproducible (every component is built from code), observable '
        '(metrics and logs flow to centralised monitoring), and resilient (blue/green slots provide instant recovery, '
        'and infrastructure can be destroyed and rebuilt from scratch). These properties constitute the foundation of '
        'a mature Continuous Delivery practice.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════
    # SECTION 13 — REFERENCES
    # ═══════════════════════════════════════════
    add_formatted_paragraph(doc, '13. References', bold=True, font_size=20,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    references = [
        'Fowler, M. and Lewis, J. (2014) Microservices: A Definition of This New Architectural Term. Available at: https://martinfowler.com/articles/microservices.html',
        'Humble, J. and Farley, D. (2010) Continuous Delivery: Reliable Software Releases through Build, Test, and Deployment Automation. Addison-Wesley Professional.',
        'IEA (2024) Data Centres and Data Transmission Networks. International Energy Agency.',
        'Morris, K. (2016) Infrastructure as Code: Managing Servers in the Cloud. O\'Reilly Media.',
        'Newman, S. (2021) Building Microservices: Designing Fine-Grained Systems. 2nd edn. O\'Reilly Media.',
        'Potvin, R. and Levenberg, J. (2016) \'Why Google Stores Billions of Lines of Code in a Single Repository\', Communications of the ACM, 59(7), pp. 78–87.',
        'Burns, B. et al. (2022) Kubernetes: Up and Running. 3rd edn. O\'Reilly Media.',
        'Microsoft (2025) Azure Kubernetes Service Documentation. Available at: https://learn.microsoft.com/en-us/azure/aks/',
        'Microsoft (2025) Azure Well-Architected Framework. Available at: https://learn.microsoft.com/en-us/azure/well-architected/',
        'HashiCorp (2025) Terraform AzureRM Provider Documentation. Available at: https://registry.terraform.io/providers/hashicorp/azurerm/latest/docs',
        'Aqua Security (2025) Trivy Documentation. Available at: https://aquasecurity.github.io/trivy/',
        'OWASP Foundation (2021) OWASP Top Ten. Available at: https://owasp.org/www-project-top-ten/',
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.text = ref
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        for run in p.runs:
            run.font.size = Pt(10)

    # ═══════════════════════════════════════════
    # APPENDICES
    # ═══════════════════════════════════════════
    doc.add_page_break()
    add_formatted_paragraph(doc, 'Appendix A: GitHub Secrets Configuration', bold=True, font_size=16,
                           color=(0x1e, 0x3a, 0x5f), space_after=8)

    add_table(doc,
        ['Secret Name', 'Description', 'Source'],
        [
            ['AZURE_CREDENTIALS', 'Azure service principal JSON', 'az ad sp create-for-rbac --sdk-auth'],
            ['ACR_NAME', 'ACR resource name', 'Terraform output acr_name'],
            ['ACR_LOGIN_SERVER', 'ACR login server URL', 'Terraform output acr_login_server'],
            ['ARM_CLIENT_ID', 'Service principal client ID', 'SP creation'],
            ['ARM_CLIENT_SECRET', 'Service principal client secret', 'SP creation'],
            ['ARM_SUBSCRIPTION_ID', 'Azure subscription ID', 'az account show'],
            ['ARM_TENANT_ID', 'Azure AD tenant ID', 'az account show'],
        ])

    add_formatted_paragraph(doc, 'Appendix B: Quick Command Reference', bold=True, font_size=16,
                           color=(0x1e, 0x3a, 0x5f), space_after=8, space_before=16)

    commands = [
        ('Local Development', [
            ('docker-compose up --build', 'Start all services locally'),
            ('docker-compose down', 'Stop all services'),
        ]),
        ('Infrastructure', [
            ('./scripts/setup.sh', 'Provision Azure infrastructure'),
            ('./scripts/destroy.sh', 'Destroy all infrastructure'),
        ]),
        ('Deployment', [
            ('./scripts/deploy.sh green', 'Deploy to green slot'),
            ('./scripts/deploy.sh blue', 'Deploy to blue slot'),
        ]),
        ('Rollback', [
            ('./scripts/rollback.sh', 'Switch traffic to previous slot'),
        ]),
        ('Monitoring', [
            ('kubectl get pods -n recipe-app', 'Check pod status'),
            ('kubectl logs -f <pod> -n recipe-app', 'Stream pod logs'),
            ('kubectl top pods -n recipe-app', 'Resource utilisation'),
        ]),
    ]
    for category, cmds in commands:
        p = doc.add_paragraph()
        run = p.add_run(category)
        run.bold = True
        for cmd, desc in cmds:
            p = doc.add_paragraph()
            run = p.add_run(cmd)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.add_run(f'  —  {desc}')
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(2)

    # ═══════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════
    output_path = os.path.join(os.path.dirname(__file__), 'Release_Management_Plan.docx')
    doc.save(output_path)
    print(f'Word document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    build_document()
